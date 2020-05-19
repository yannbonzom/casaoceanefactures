from docx import Document
import math
from docx.shared import Pt
import sqlite3
from os import path

# Root for the db so that it works locally and on PythonAnywhere
ROOT = path.dirname(path.realpath(__file__))

# Generate client receipt
def clientReceipt(source, firstName, lastName, sex, startDate, endDate, lengthOfStay, receiptID, totalPrice):
    # Before proceeding, store all data in the db
    connection = sqlite3.connect(path.join(ROOT, "receipts.db"))
    cursor = connection.cursor()
    cursor.execute("INSERT INTO receipts (type, source, firstName, lastName, sex, startDate, endDate, lengthOfStay, receiptID, totalPrice) VALUES ('client', ?, ?, ?, ?, ?, ?, ?, ?, ?)", (source, firstName, lastName, sex, startDate, endDate, lengthOfStay, receiptID, totalPrice))
    cursor.close()
    connection.commit()
    connection.close()

    # Price calculations
    withoutTax = round_up(totalPrice / 1.2, 2)
    tax = round_down(0.2 * withoutTax, 2)

    # Open template which we'll later save in another document
    document = Document(path.join(ROOT, 'clientTemplate.docx'))

    # Font settings
    # For all 'paragraphs'
    font = document.styles['Normal'].font
    font.name = 'Arial'
    font.size = Pt(13)
    font.bold = True
    # For 'Heading 1'
    font = document.styles['Heading 1'].font
    font.name = 'Arial'
    font.size = Pt(16)
    font.bold = True
    # For 'Heading 2'
    font = document.styles['Heading 2'].font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = False

    # Populate document with user-data 
    # In paragraphs
    for paragraph in document.paragraphs:
        if "<firstParagraph>" in paragraph.text:
            if lengthOfStay == 1:
                if startDate.year != endDate.year:
                    paragraph.text = f"Durée séjour de {sex} {lastName}: 1 jour, du {two(startDate.day)}/{two(startDate.month)}/{startDate.year} au {two(endDate.day)}/{two(endDate.month)}/{endDate.year},"
                else: 
                    paragraph.text = f"Durée séjour de {sex} {lastName}: 1 jour, du {two(startDate.day)}/{two(startDate.month)} au {two(endDate.day)}/{two(endDate.month)}/{endDate.year},"
            else: 
                if startDate.year != endDate.year:
                    paragraph.text = f"Durée séjour de {sex} {lastName}: {lengthOfStay} jours, du {two(startDate.day)}/{two(startDate.month)}/{startDate.year} au {two(endDate.day)}/{two(endDate.month)}/{endDate.year},"
                else: 
                    paragraph.text = f"Durée séjour de {sex} {lastName}: {lengthOfStay} jours, du {two(startDate.day)}/{two(startDate.month)} au {two(endDate.day)}/{two(endDate.month)}/{endDate.year},"
        elif "<firstParagraphPricePart>" in paragraph.text:
            paragraph.text = f"au prix de{mad(totalPrice)} TTC."
        elif "<totalPrice>" in paragraph.text:
            paragraph.text = f"NET A PAYER EN MAD:  {mad(totalPrice)} TTC"
        elif "<inWords>" in paragraph.text:
            paragraph.text = f"{toWords(totalPrice)}, TTC"
    # In top table
    document.tables[0].cell(1,0).text = f"N° de FACTURE {receiptID}-{endDate.year}"
    document.tables[0].cell(1,1).text = f"{sex} {firstName} {lastName}"
    document.tables[0].cell(2,1).text = f"DATE: {two(endDate.day)}/{two(endDate.month)}/{endDate.year}"
    # In center table (with left-alignment)
    number1 = document.tables[1].cell(0,1).paragraphs[0]
    number1.add_run(f"{mad(withoutTax)}")
    number2 = document.tables[2].cell(0,1).paragraphs[0]
    number2.add_run(f"{mad(tax)}")
    number3 = document.tables[3].cell(0,1).paragraphs[0]
    number3.add_run(f"{mad(totalPrice)}")

    # Save document with formatted name
    docName = "FACTURE MDB Maroc Services " + str(receiptID) + "-" + str(endDate.year) + ".docx"
    document.save(path.join(ROOT, docName))
    
    # Download it to the user's computer
    return docName

# Generate seminar receipt
def seminarReceipt(company, receiptID, startDate, endDate, lengthOfStay, pricePerDayWithoutTax):
    # Before proceeding, store all data in the db
    connection = sqlite3.connect(path.join(ROOT, "receipts.db"))
    cursor = connection.cursor()
    cursor.execute("INSERT INTO receipts (type, company, receiptID, startDate, endDate, lengthOfStay, pricePerDayWithoutTax) VALUES ('seminar', ?, ?, ?, ?, ?, ?)", (company, receiptID, startDate, endDate, lengthOfStay, pricePerDayWithoutTax))
    cursor.close()
    connection.commit()
    connection.close()

    # Price calculations
    withoutTax = float(lengthOfStay) * pricePerDayWithoutTax
    tax = round_down(0.2 * withoutTax)
    totalPrice = round_up(withoutTax + tax)

    # Open template which we'll later save in another document
    document = Document(path.join(ROOT, 'seminarTemplate.docx'))

    # Font settings
    # For all 'paragraphs'
    font = document.styles['Normal'].font
    font.name = 'Arial'
    font.size = Pt(13)
    font.bold = True
    # For 'Heading 1'
    font = document.styles['Heading 1'].font
    font.name = 'Arial'
    font.size = Pt(16)
    font.bold = True
    # For 'Heading 2'
    font = document.styles['Heading 2'].font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = False

    # Populate document with user-data 
    # In paragraphs
    for paragraph in document.paragraphs:
        if "<firstParagraph>" in paragraph.text:
            if lengthOfStay == 1:
                paragraph.text = f"SEMINAIRE: 1 jour, le {two(startDate.day)}/{two(startDate.month)}/{startDate.year}, au prix de"
            else: 
                if startDate.year != endDate.year:
                    paragraph.text = f"SEMINAIRE: {lengthOfStay} jours, du {two(startDate.day)}/{two(startDate.month)}/{startDate.year} au {two(endDate.day)}/{two(endDate.month)}/{endDate.year}, au prix par jour"
                else: 
                    paragraph.text = f"SEMINAIRE: {lengthOfStay} jours, du {two(startDate.day)}/{two(startDate.month)} au {two(endDate.day)}/{two(endDate.month)}/{endDate.year}, au prix par jour"
        elif "<firstParagraphPricePart>" in paragraph.text:
            paragraph.text = f"de{mad(pricePerDayWithoutTax)} HT."
        elif "<totalPrice>" in paragraph.text:
            paragraph.text = f"NET A PAYER EN MAD:  {mad(totalPrice)} TTC"
        elif "<inWords>" in paragraph.text:
            paragraph.text = f"{toWords(totalPrice)}, TTC"
    # In top table
    document.tables[0].cell(1,0).text = f"N° de FACTURE {receiptID}-{endDate.year}"
    document.tables[0].cell(2,1).text = f"DATE: {two(endDate.day)}/{two(endDate.month)}/{endDate.year}"
    # In center table (with left-alignment)
    number1 = document.tables[1].cell(0,1).paragraphs[0]
    number1.add_run(f"{mad(withoutTax)}")
    number2 = document.tables[2].cell(0,1).paragraphs[0]
    number2.add_run(f"{mad(tax)}")
    number3 = document.tables[3].cell(0,1).paragraphs[0]
    number3.add_run(f"{mad(totalPrice)}")

    # Save document with formatted name
    docName = "FACTURE MDB Maroc Services " + str(receiptID) + "-" + str(endDate.year) + " " + company + ".docx"
    document.save(path.join(ROOT, docName))
    
    # Download it to the user's computer
    return docName


# FUNCTIONS NECESSARY FOR clientReceipt() and seminarReceipt()
# Format money values in MAD
def mad(n):
    formatted = '{: ,.2f}'.format(n).replace(',', ' ').replace('.', ',')
    return f"{formatted} MAD"
# Format dates to two digits (ex. 1 -> 01)
def two(n):
    formatted = '{0:02}'.format(n)
    return formatted
# Rounding definitions
def round_up(n, decimals=0):
    multiplier = 10 ** decimals
    return math.ceil(n * multiplier) / multiplier
def round_down(n, decimals=0):
    multiplier = 10 ** decimals
    return math.floor(n * multiplier) / multiplier
# Converting price in numbers into words
def toWords(n):
    # inWords string
    inWords = ""

    # Convert number into formatted string for easy manipulation
    n = f"{n: ,.2f}"

    # Split n into whole and decimal parts based on period
    nSplitted = n.split(".")
    whole = nSplitted[0]
    decimal = nSplitted[1]

    # Split whole part into blocks of three numbers based on commas
    threes = whole.split(",")

    # Iterate over each block of threes to write it in words
    # If in the thousands
    if len(threes) == 2:
        inWords = read(threes[0]) + " mille " + read(threes[1]) + " dirhams"
    else:
        inWords = read(threes[0]) + " dirhams"

    # Read the cents
    if decimal == "01":
        inWords += " et un centime"
    elif decimal != "00":
        inWords += " et " + read(decimal) + " centimes"
    
    # Remove 'un' if it appears at the string's beginning
    unFix = inWords.split()
    if unFix[0] == "un":
        i = 1
        inWords = ""
        while i < len(unFix):
            # if/else accounts for space added to inWords's very end
            if i != len(unFix) - 1:
                inWords += unFix[i] + " "
            else:
                inWords += unFix[i]
            i += 1
    
    # Return capitalized string
    return inWords.capitalize()
# Part of the toWords function
def read(threes):
    # Set up number dictionaries
    numbers = {
        "1": "un",
        "2": "deux",
        "3": "trois",
        "4": "quatre",
        "5": "cinq",
        "6": "six",
        "7": "sept",
        "8": "huit", 
        "9": "neuf", 
        "11": "onze",
        "12": "douze", 
        "13": "treize",
        "14": "quatorze",
        "15": "quinze",
        "16": "seize",
        "71": "soixante et onze",
        "72": "soixante-douze",
        "73": "soixante-treize",
        "74": "soixante-quatorze",
        "75": "soixante-quinze",
        "76": "soixante-seize",
        "91": "quatre-vingt-onze",
        "92": "quatre-vingt-douze",
        "93": "quatre-vingt-treize",
        "94": "quatre-vingt-quatorze",
        "95": "quatre-vingt-quinze",
        "96": "quatre-vingt-seize",
    }
    tens = {
        "1": "dix", 
        "2": "vingt",
        "3": "trente",
        "4": "quarante", 
        "5": "cinquante",
        "6": "soixante",
        "7": "soixante-dix",
        "8": "quatre-vingt",
        "9": "quatre-vingt-dix",
    }

    # Make any block of threes actually have three numbers (ex. 9 --> 009)
    threes = int(threes)
    threes = "{0:03}".format(threes)

    # Declare the return value toWords
    toWords = ""
    
    # State hundreds place
    if threes[0] != "0":
        # Account for spelling difference cent-cents
        if threes[1] == "0" and threes[2] == "0":
            toWords = numbers[threes[0]] + " cents"
            return toWords
        else:
            toWords = numbers[threes[0]] + " cent "
    # Account for numbers from 11-16
    if f"{threes[1]}{threes[2]}" in numbers.keys():
        toWords += numbers[f"{threes[1]}{threes[2]}"]
        return toWords
    # State tens place (if-elif accounts for correct spacing between the number and 'dirhams')
    if threes[1] != "0" and threes[2] != "0":
        # Account for '-' connection when in 70s or 90s
        if threes[1] == "7" or threes[1] == "9":
            toWords += tens[threes[1]]
        else: 
            toWords += tens[threes[1]] + " "
    elif threes[1] != "0":
        # Account for '-' connection when in 70s or 90s
        if threes[1] == "7" or threes[1] == "9":
            toWords += tens[threes[1]]
        else: 
            toWords += tens[threes[1]]
    # State ones place
    if threes[2] != "0":
        # Account for '-' connection when in 70s or 90s
        if threes[1] == "7" or threes[1] == "9":
            toWords += "-" + numbers[threes[2]]
        else: 
            toWords += numbers[threes[2]]

    # Return the written out form
    return toWords